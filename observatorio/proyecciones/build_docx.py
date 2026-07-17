#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
FUTURIA · Generador de DOCX con marca de agua nativa de Word + proteccion Office.

- Marca de agua diagonal "FUTURIA · CONFIDENCIAL" en TODAS las paginas (header con
  imagen PNG rotada + wordart-like; indeleble al imprimir / exportar a PDF).
- Proteccion: documento de SOLO LECTURA con contrasena de modificacion (settings.xml).
  Al abrir pide contrasena para editar; para leer/imprimir no.
- Portada con trazabilidad (titulo, control, fecha).

Uso:
  python build_docx.py entrada.md --out salida.docx \
      --titulo "Proyeccion Sectorial: Grupo SanCor Salud" \
      --control "FUT-SANCOR-2026-001" \
      --watermark "FUTURIA · CONFIDENCIAL" \
      --owner "FUTURIA-LOCK-2026"
"""
import argparse, os, datetime, zipfile, re
import markdown
from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH

def generar_png_marca(texto, path, angulo=-30):
    ancho, alto = 1100, 1500
    img = Image.new("RGBA", (ancho, alto), (255, 255, 255, 0))
    from PIL import ImageDraw
    d = ImageDraw.Draw(img)
    try:
        font = ImageFont.truetype("arial.ttf", 46)
    except Exception:
        font = ImageFont.load_default()
    step_x, step_y = 470, 300
    for y in range(-alto, alto, step_y):
        for x in range(-ancho, ancho, step_x):
            d.text((x, y), texto, font=font, fill=(130, 130, 140, 60))
    img = img.rotate(angulo, expand=False, center=(ancho // 2, alto // 2))
    img.save(path, "PNG")
    return path

def agregar_marca_de_agua(doc, png_path, texto):
    """Inserta la imagen de marca de agua en el header (primario) de forma diagonal."""
    header = doc.sections[0].header
    # asegurar que el header tenga un parrafo
    if not header.paragraphs:
        header.add_paragraph()
    para = header.paragraphs[0]
    run = para.add_run()
    # tamaño de página en emu (~ A4)
    from docx.shared import Emu
    run.add_picture(png_path, width=Emu(19000000), height=Emu(25000000))
    # posicion absoluta centrada, detras del texto, rotada
    drawing = run._r.find(qn('w:drawing'))
    if drawing is not None:
        inline = drawing.find(qn('wp:inline'))
        if inline is not None:
            # convertir a anchor (posicion absoluta)
            extent = inline.find(qn('wp:extent'))
            docPr = inline.find(qn('wp:docPr'))
            # construir anchor
            anchor = OxmlElement('wp:anchor')
            anchor.set('distT', '0'); anchor.set('distB', '0')
            anchor.set('distL', '0'); anchor.set('distR', '0')
            anchor.set('simplePos', '0'); anchor.set('relativeHeight', '1')
            anchor.set('behindDoc', '1'); anchor.set('locked', '0')
            anchor.set('layoutInCell', '1'); anchor.set('allowOverlap', '1')
            simple = OxmlElement('wp:simplePos'); simple.set('x', '0'); simple.set('y', '0')
            posH = OxmlElement('wp:positionH'); posH.set('relativeFrom', 'page')
            ah = OxmlElement('wp:align'); ah.text = 'center'
            posH.append(ah)
            posV = OxmlElement('wp:positionV'); posV.set('relativeFrom', 'page')
            av = OxmlElement('wp:align'); av.text = 'center'
            posV.append(av)
            anchor.append(simple); anchor.append(posH); anchor.append(posV)
            anchor.append(extent); anchor.append(docPr)
            for child in list(inline):
                if child.tag in (qn('wp:graphic'), qn('a:graphic')):
                    anchor.append(child)
            drawing.append(anchor)
            drawing.remove(inline)
    # repetir en header primario para todas las paginas (ya aplica a todas las secciones)

def proteger(docx_path, owner_password):
    """Setea documento de solo lectura con contrasena de modificacion (settings.xml)."""
    import hashlib
    # Word usa hash + salt del password para modificar
    salt = os.urandom(16)
    def hash_password(pwd, salt):
        h = hashlib.sha1(salt + pwd.encode('utf-8')).digest()
        for _ in range(49999):
            h = hashlib.sha1(h).digest()
        return h
    hp = hash_password(owner_password, salt).hex()
    salt_hex = salt.hex()
    with zipfile.ZipFile(docx_path, 'r') as z:
        names = z.namelist()
        data = {n: z.read(n) for n in names}
    settings = data.get('word/settings.xml', b'<w:settings xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>')
    settings = settings.decode('utf-8')
    if '<w:documentProtection' in settings:
        settings = re.sub(r'<w:documentProtection[^>]*/?>', '', settings)
    tag = ('<w:documentProtection w:edit="readOnly" '
           'w:enforcement="1" w:cookieProtected="1" '
           'w:hash="{hp}" w:salt="{salt}" w:algorithmSid="SHA-1"/>').format(hp=hp, salt=salt_hex)
    # insertar antes de </w:settings>
    settings = settings.replace('</w:settings>', tag + '</w:settings>')
    data['word/settings.xml'] = settings.encode('utf-8')
    with zipfile.ZipFile(docx_path, 'w', zipfile.ZIP_DEFLATED) as z:
        for n, b in data.items():
            z.writestr(n, b)

def md_a_docx(md_path, out_path, titulo, control, watermark="FUTURIA · CONFIDENCIAL",
              owner="FUTURIA-LOCK-2026", logo=None, sello="BORRADOR — NO VERIFICADO",
              audit_md=None, traz_pct=None):
    base = os.path.dirname(os.path.abspath(out_path))
    png = os.path.join(base, "_wm_tmp.png")
    generar_png_marca(watermark, png)

    doc = Document()
    # estilos base
    normal = doc.styles['Normal']
    normal.font.name = 'Calibri'
    normal.font.size = Pt(11)

    # PORTADA
    for _ in range(2):
        doc.add_paragraph()
    # logo (placeholder o real)
    if logo and os.path.exists(logo):
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(logo, width=Inches(1.6))
    else:
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run("▲ FUTURIA"); r.bold = True; r.font.size = Pt(20)
        r.font.color.rgb = RGBColor(0xF5, 0x50, 0x32)
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run("«LOGO PENDIENTE DE ARCHIVO OFICIAL»"); r.italic = True; r.font.size = Pt(8)
        r.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("FUTURIA · INSTITUTO GLOBAL DE FUTUROS PLAUSIBLES"); r.bold = True; r.font.size = Pt(11)
    r.font.color.rgb = RGBColor(0x88, 0x88, 0x99)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(titulo); r.bold = True; r.font.size = Pt(20)
    r.font.color.rgb = RGBColor(0xF5, 0x50, 0x32)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Observatorio de Vigilancia Estrategica\nProspeccion de Negocios con trazabilidad"); r.font.size = Pt(11)
    r.font.color.rgb = RGBColor(0x44, 0x44, 0x55)
    # sello SAT
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("⚠ " + sello); r.bold = True; r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(0xB0, 0x30, 0x20)
    if traz_pct is not None:
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run("Trazabilidad: %s de hechos con origen verificable" % traz_pct); r.italic = True; r.font.size = Pt(9)
        r.font.color.rgb = RGBColor(0x88, 0x88, 0x99)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Fecha: %s    |    Control: %s" % (datetime.date.today().isoformat(), control))
    r.italic = True; r.font.size = Pt(9); r.font.color.rgb = RGBColor(0x88, 0x88, 0x99)
    doc.add_page_break()

    # marcar agua en header (todas las paginas)
    agregar_marca_de_agua(doc, png, watermark)

    # cuerpo desde markdown
    html = markdown.markdown(open(md_path, encoding='utf-8').read(),
                             extensions=['tables', 'fenced_code'])
    # parsear HTML simple: headings, p, table, lists
    import xml.etree.ElementTree as ET
    html = html.replace('&nbsp;', ' ')
    # envolver en root
    root = ET.fromstring('<root>' + html + '</root>')
    def add_runs(parent_p, elem):
        if elem.text:
            parent_p.add_run(elem.text)
        for c in elem:
            if c.tag in ('strong','b'):
                run = parent_p.add_run(c.text or '')
                run.bold = True
            elif c.tag in ('em','i'):
                run = parent_p.add_run(c.text or '')
                run.italic = True
            else:
                if c.text:
                    parent_p.add_run(c.text)
            if c.tail:
                parent_p.add_run(c.tail)
    def walk(elem):
        for child in elem:
            tag = child.tag
            if tag in ('h1','h2','h3','h4'):
                lvl = int(tag[1])
                p = doc.add_heading(level=lvl)
                add_runs(p, child)
            elif tag == 'p':
                p = doc.add_paragraph()
                add_runs(p, child)
            elif tag == 'ul':
                for li in child:
                    p = doc.add_paragraph(style='List Bullet')
                    add_runs(p, li)
            elif tag == 'ol':
                for li in child:
                    p = doc.add_paragraph(style='List Number')
                    add_runs(p, li)
            elif tag == 'table':
                rows = child.findall('tr')
                if rows:
                    ncol = max(len(r.findall('td')) for r in rows)
                    t = doc.add_table(rows=len(rows), cols=ncol)
                    t.style = 'Light Grid Accent 1'
                    for i, r in enumerate(rows):
                        cells = r.findall('td') + r.findall('th')
                        for j, c in enumerate(cells):
                            if j < ncol:
                                t.rows[i].cells[j].text = (c.text or '').strip()
            elif tag in ('blockquote','div'):
                walk(child)
    walk(root)

    # anexo de auditoria SAT (opcional)
    if audit_md and os.path.exists(audit_md):
        doc.add_page_break()
        doc.add_heading("Anexo: Auditoria de Trazabilidad (SAT-FUTURIA)", level=1)
        for line in open(audit_md, encoding='utf-8').read().splitlines():
            line = line.rstrip()
            if not line:
                continue
            if line.startswith('#'):
                lvl = min(len(line) - len(line.lstrip('#')), 4)
                doc.add_heading(line.lstrip('#').strip(), level=lvl)
            else:
                doc.add_paragraph(line)

    doc.save(out_path)
    try:
        os.remove(png)
    except Exception:
        pass
    # proteccion SOLO si el usuario la pide explicitamente con --owner
    if owner:
        proteger(out_path, owner)
    return out_path

if __name__ == "__main__":
    ap = argparse.ArgumentParser()
    ap.add_argument("entrada")
    ap.add_argument("--out", default="salida.docx")
    ap.add_argument("--titulo", default="Informe FUTURIA")
    ap.add_argument("--control", default="FUT-CTRL")
    ap.add_argument("--watermark", default="FUTURIA · CONFIDENCIAL")
    ap.add_argument("--owner", default=None, help="opcional: contrasena de modificacion (solo lectura). Si se omite, el doc NO se protege.")
    ap.add_argument("--logo", default=None, help="ruta al logo FUTURIA (png/jpg)")
    ap.add_argument("--sello", default="BORRADOR — NO VERIFICADO", help="sello SAT en portada")
    ap.add_argument("--audit", default=None, help="ruta a .md de auditoria SAT a incluir como anexo")
    ap.add_argument("--traz", default=None, help="texto % trazabilidad, ej. '14/14'")
    a = ap.parse_args()
    ruta = md_a_docx(a.entrada, a.out, a.titulo, a.control, a.watermark, a.owner,
                     logo=a.logo, sello=a.sello, audit_md=a.audit, traz_pct=a.traz)
    print("DOCX generado:", ruta)
